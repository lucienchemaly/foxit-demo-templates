"""
Build signing_agreement.docx, a Microsoft Word starting point for a Foxit eSign
template. The three eSign Text Tags below map to party 1 and become interactive
fields when the document is uploaded to the eSign editor or sent through the
templates / createfolder API with "processTextTags": true.

Tags (see Foxit eSign Text Tags docs):
    ${textfield:1:y:Name_of_signer:______}  -> signer name (required text field)
    ${datefield:1:n::____}                   -> date field
    ${s:1:______}                            -> signature field

Placeholders must use underscores; an empty placeholder like ${s:1: } does not
produce a signature field and the resulting template cannot be sent.

Run:
    pip install python-docx
    python3 build_signing_agreement.py
"""
from docx import Document

doc = Document()
doc.add_heading("Service Agreement", level=1)
doc.add_paragraph(
    "This agreement is provided as a sample template for the Foxit eSign API "
    "walkthrough. By signing below, the signer accepts the terms of service."
)
doc.add_paragraph("Full name: ${textfield:1:y:Name_of_signer:______}")
doc.add_paragraph("Date: ${datefield:1:n::____}")
doc.add_paragraph("Signature: ${s:1:______}")
doc.save("signing_agreement.docx")
print("wrote signing_agreement.docx")
