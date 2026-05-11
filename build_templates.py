"""
Build the Word templates referenced by article1-writing.md / article1-client.md /
article2-batch3-writing.md and verify each one against the Foxit DocGen API.

Outputs:
  - invoice_simple.docx          (scalar tokens only, first-run friendly)
  - invoice_table.docx           (full invoice with line-items loop, camelCase tokens)
  - invoice_full.docx            (snake_case invoice, matches article1-client.md)
  - contract_standard.docx       (two-party MSA, fixed term, batch3 article2)
  - contract_auto_renewal.docx   (two-party MSA, auto-renewing, batch3 article2)
  - compliance_attestation.docx  (quarterly vendor risk attestation, batch3 article2)
  - *_test.pdf rendered proof for each template
"""
import base64
import os
import sys
from pathlib import Path

import requests
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Inches

HERE = Path(__file__).resolve().parent

# Demo credentials are documented in foxit/CLAUDE.md; pulled from env first
# so the script can run in CI without a code change.
HOST = os.environ.get("BASE_URL", "https://na1.fusion.foxit.com")
CLIENT_ID = os.environ.get("CLIENT_ID", "foxit_8Aqfv2MCkou5rr4i")
CLIENT_SECRET = os.environ.get("CLIENT_SECRET", "_ySDdB1dKLP3FVjkLmttgIG1X1B_oHGZ")


def _styled_heading(doc, text, size=18):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return p


def build_invoice_simple(path: Path) -> None:
    """Scalar-only invoice header. Good for first end-to-end run."""
    doc = Document()

    _styled_heading(doc, "INVOICE", size=24)
    doc.add_paragraph("Issued by Foxit DocGen API Demo")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Bill To: ").bold = True
    p.add_run("{{ companyName }}")

    p = doc.add_paragraph()
    p.add_run("Invoice Number: ").bold = True
    p.add_run("{{ invoiceNumber }}")

    p = doc.add_paragraph()
    p.add_run("Invoice Date: ").bold = True
    p.add_run("{{ invoiceDate \\@ MM/dd/yyyy }}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Total Due: ").bold = True
    run = p.add_run('{{ totalDue \\# "$#,##0.00" }}')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Thank you for your business. Payment is due within 30 days of the invoice date."
    )

    doc.save(path)


def build_invoice_table(path: Path) -> None:
    """Full invoice with line-items loop. Mirrors the article's JSON payload exactly."""
    doc = Document()

    _styled_heading(doc, "INVOICE", size=24)
    doc.add_paragraph("Foxit DocGen API Sample Template")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Bill To: ").bold = True
    p.add_run("{{ companyName }}")

    p = doc.add_paragraph()
    p.add_run("Invoice Number: ").bold = True
    p.add_run("{{ invoiceNumber }}")

    p = doc.add_paragraph()
    p.add_run("Invoice Date: ").bold = True
    p.add_run("{{ invoiceDate \\@ MM/dd/yyyy }}")

    doc.add_paragraph()

    # Line-items table: header row (static) + loop row (TableStart/TableEnd in same row)
    # + footer row using {{=SUM(ABOVE)}} aggregate.
    table = doc.add_table(rows=3, cols=5)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    headers = ["#", "Description", "Qty", "Unit Price", "Line Total"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    loop_row = table.rows[1].cells
    loop_row[0].text = "{{TableStart:lineItems}}{{ROW_NUMBER}}"
    loop_row[1].text = "{{description}}"
    loop_row[2].text = "{{qty}}"
    loop_row[3].text = '{{unitPrice \\# "$#,##0.00"}}'
    loop_row[4].text = '{{lineTotal \\# "$#,##0.00"}}{{TableEnd:lineItems}}'

    footer = table.rows[2].cells
    footer[0].text = ""
    footer[1].text = ""
    footer[2].text = ""
    p = footer[3].paragraphs[0]
    p.add_run("Subtotal:").bold = True
    footer[4].text = '{{=SUM(ABOVE) \\# "$#,##0.00"}}'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Total Due: ").bold = True
    p.add_run('{{ totalDue \\# "$#,##0.00" }}').bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Thank you for your business. Payment is due within 30 days of the invoice date."
    )

    doc.save(path)


def build_invoice_full(path: Path) -> None:
    """Snake_case invoice that matches article1-client.md's document_values exactly."""
    doc = Document()

    _styled_heading(doc, "INVOICE", size=24)
    doc.add_paragraph("Foxit DocGen API Sample Template")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Bill To: ").bold = True
    p.add_run("{{ customer_name }}")

    p = doc.add_paragraph()
    p.add_run("Invoice Number: ").bold = True
    p.add_run("{{ invoice_number }}")

    p = doc.add_paragraph()
    p.add_run("Invoice Date: ").bold = True
    p.add_run("{{ invoice_date }}")

    p = doc.add_paragraph()
    p.add_run("Due Date: ").bold = True
    p.add_run("{{ due_date }}")

    doc.add_paragraph()

    # Line-items table: header + loop row + subtotal/tax/total footer rows.
    table = doc.add_table(rows=6, cols=5)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    headers = ["#", "Description", "Qty", "Unit Price", "Total"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    loop_row = table.rows[1].cells
    loop_row[0].text = "{{TableStart:line_items}}{{ROW_NUMBER}}"
    loop_row[1].text = "{{description}}"
    loop_row[2].text = "{{qty}}"
    loop_row[3].text = '{{unit_price \\# "$#,##0.00"}}'
    loop_row[4].text = '{{total \\# "$#,##0.00"}}{{TableEnd:line_items}}'

    # Footer rows: subtotal, tax, total due
    subtotal_row = table.rows[2].cells
    subtotal_row[0].text = ""
    subtotal_row[1].text = ""
    subtotal_row[2].text = ""
    subtotal_row[3].paragraphs[0].add_run("Subtotal:").bold = True
    subtotal_row[4].text = '{{subtotal \\# "$#,##0.00"}}'

    tax_row = table.rows[3].cells
    tax_row[0].text = ""
    tax_row[1].text = ""
    tax_row[2].text = ""
    tax_row[3].paragraphs[0].add_run("Tax Rate:").bold = True
    # Foxit's \# "0.00%" switch does not multiply by 100, so we render
    # tax_rate as a raw decimal (0.08). Articles can still teach the switch
    # against numeric currency fields where the multiply-by-100 is irrelevant.
    tax_row[4].text = "{{tax_rate}}"

    tax_amount_row = table.rows[4].cells
    tax_amount_row[0].text = ""
    tax_amount_row[1].text = ""
    tax_amount_row[2].text = ""
    tax_amount_row[3].paragraphs[0].add_run("Tax Amount:").bold = True
    tax_amount_row[4].text = '{{tax_amount \\# "$#,##0.00"}}'

    total_row = table.rows[5].cells
    total_row[0].text = ""
    total_row[1].text = ""
    total_row[2].text = ""
    total_row[3].paragraphs[0].add_run("Total Due:").bold = True
    p = total_row[4].paragraphs[0]
    run = p.add_run('{{total_due \\# "$#,##0.00"}}')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Thank you for your business. Payment is due within 30 days of the invoice date."
    )

    doc.save(path)


def build_contract_standard(path: Path) -> None:
    """Two-party master services agreement, fixed-term, batch3 article2."""
    doc = Document()

    _styled_heading(doc, "MASTER SERVICES AGREEMENT", size=20)
    doc.add_paragraph(
        "This Master Services Agreement (the Agreement) is entered into as of "
        "{{ effective_date }} by and between the parties listed below."
    )
    doc.add_paragraph()

    _styled_heading(doc, "Parties", size=14)

    p = doc.add_paragraph()
    p.add_run("Party A: ").bold = True
    p.add_run("{{ party_a_name }}")

    p = doc.add_paragraph()
    p.add_run("Address: ").bold = True
    p.add_run("{{ party_a_address_line1 }}, {{ party_a_address_city }}, {{ party_a_address_state }}")

    p = doc.add_paragraph()
    p.add_run("Signatory: ").bold = True
    p.add_run("{{ party_a_signatory_name }}, {{ party_a_signatory_title }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Party B: ").bold = True
    p.add_run("{{ party_b_name }}")

    p = doc.add_paragraph()
    p.add_run("Address: ").bold = True
    p.add_run("{{ party_b_address_line1 }}, {{ party_b_address_city }}, {{ party_b_address_state }}")

    p = doc.add_paragraph()
    p.add_run("Signatory: ").bold = True
    p.add_run("{{ party_b_signatory_name }}, {{ party_b_signatory_title }}")

    doc.add_paragraph()

    _styled_heading(doc, "Term and Payment", size=14)

    p = doc.add_paragraph()
    p.add_run("Effective Date: ").bold = True
    p.add_run("{{ effective_date }}")

    p = doc.add_paragraph()
    p.add_run("Term: ").bold = True
    p.add_run("{{ term_months }} months (fixed term, no automatic renewal)")

    p = doc.add_paragraph()
    p.add_run("Contract Value: ").bold = True
    p.add_run('{{ contract_value \\# "$#,##0.00" }}')

    p = doc.add_paragraph()
    p.add_run("Payment Schedule: ").bold = True
    p.add_run("{{ payment_schedule }}")

    p = doc.add_paragraph()
    p.add_run("Governing Law: ").bold = True
    p.add_run("{{ governing_law }}")

    doc.add_paragraph()

    _styled_heading(doc, "Termination", size=14)
    doc.add_paragraph(
        "This Agreement expires at the end of the term stated above. The parties may "
        "extend the term only by signing a new written agreement."
    )

    doc.add_paragraph()
    _styled_heading(doc, "Signatures", size=14)
    doc.add_paragraph("Party A: ________________________________  Date: __________")
    doc.add_paragraph("Party B: ________________________________  Date: __________")

    doc.save(path)


def build_contract_auto_renewal(path: Path) -> None:
    """Two-party master services agreement, auto-renewing, batch3 article2."""
    doc = Document()

    _styled_heading(doc, "MASTER SERVICES AGREEMENT (Auto-Renewing)", size=20)
    doc.add_paragraph(
        "This Master Services Agreement (the Agreement) is entered into as of "
        "{{ effective_date }} by and between the parties listed below."
    )
    doc.add_paragraph()

    _styled_heading(doc, "Parties", size=14)

    p = doc.add_paragraph()
    p.add_run("Party A: ").bold = True
    p.add_run("{{ party_a_name }}")

    p = doc.add_paragraph()
    p.add_run("Address: ").bold = True
    p.add_run("{{ party_a_address_line1 }}, {{ party_a_address_city }}, {{ party_a_address_state }}")

    p = doc.add_paragraph()
    p.add_run("Signatory: ").bold = True
    p.add_run("{{ party_a_signatory_name }}, {{ party_a_signatory_title }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Party B: ").bold = True
    p.add_run("{{ party_b_name }}")

    p = doc.add_paragraph()
    p.add_run("Address: ").bold = True
    p.add_run("{{ party_b_address_line1 }}, {{ party_b_address_city }}, {{ party_b_address_state }}")

    p = doc.add_paragraph()
    p.add_run("Signatory: ").bold = True
    p.add_run("{{ party_b_signatory_name }}, {{ party_b_signatory_title }}")

    doc.add_paragraph()

    _styled_heading(doc, "Term and Payment", size=14)

    p = doc.add_paragraph()
    p.add_run("Effective Date: ").bold = True
    p.add_run("{{ effective_date }}")

    p = doc.add_paragraph()
    p.add_run("Initial Term: ").bold = True
    p.add_run("{{ term_months }} months")

    p = doc.add_paragraph()
    p.add_run("Contract Value: ").bold = True
    p.add_run('{{ contract_value \\# "$#,##0.00" }}')

    p = doc.add_paragraph()
    p.add_run("Payment Schedule: ").bold = True
    p.add_run("{{ payment_schedule }}")

    p = doc.add_paragraph()
    p.add_run("Governing Law: ").bold = True
    p.add_run("{{ governing_law }}")

    doc.add_paragraph()

    _styled_heading(doc, "Auto-Renewal", size=14)
    doc.add_paragraph(
        "This Agreement renews automatically for successive terms equal in length to the "
        "Initial Term unless either party provides written notice of non-renewal at least "
        "thirty (30) days before the end of the then-current term. Renewal terms inherit "
        "all provisions of this Agreement unchanged."
    )

    doc.add_paragraph()
    _styled_heading(doc, "Signatures", size=14)
    doc.add_paragraph("Party A: ________________________________  Date: __________")
    doc.add_paragraph("Party B: ________________________________  Date: __________")

    doc.save(path)


def build_compliance_attestation(path: Path) -> None:
    """Quarterly vendor risk attestation report, batch3 article2."""
    doc = Document()

    _styled_heading(doc, "QUARTERLY VENDOR RISK ATTESTATION", size=20)
    doc.add_paragraph("Foxit DocGen API Sample Template")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Organization: ").bold = True
    p.add_run("{{ organizationName }}")

    p = doc.add_paragraph()
    p.add_run("Attestation Period: ").bold = True
    p.add_run("{{ attestationPeriod }}")

    p = doc.add_paragraph()
    p.add_run("Report Date: ").bold = True
    p.add_run("{{ reportDate }}")

    p = doc.add_paragraph()
    p.add_run("Prepared By: ").bold = True
    p.add_run("{{ preparedBy }}, {{ preparedByTitle }}")

    p = doc.add_paragraph()
    p.add_run("Policy Version: ").bold = True
    p.add_run("{{ policyVersion }}")

    doc.add_paragraph()

    _styled_heading(doc, "Controls Reviewed", size=14)

    table = doc.add_table(rows=2, cols=5)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    headers = ["#", "Control ID", "Control Name", "Status", "Evidence"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    loop_row = table.rows[1].cells
    loop_row[0].text = "{{TableStart:controls}}{{ROW_NUMBER}}"
    loop_row[1].text = "{{control_id}}"
    loop_row[2].text = "{{control_name}}"
    loop_row[3].text = "{{status}}"
    loop_row[4].text = "{{evidenceLink}}{{TableEnd:controls}}"

    doc.add_paragraph()

    _styled_heading(doc, "Summary", size=14)

    p = doc.add_paragraph()
    p.add_run("Total Controls Reviewed: ").bold = True
    p.add_run("{{ totalControls }}")

    p = doc.add_paragraph()
    p.add_run("Compliant: ").bold = True
    p.add_run("{{ compliantCount }}")

    p = doc.add_paragraph()
    p.add_run("In Remediation: ").bold = True
    p.add_run("{{ remediationCount }}")

    doc.add_paragraph()
    _styled_heading(doc, "Approval", size=14)
    doc.add_paragraph(
        "I attest that the information presented in this report is accurate to the best "
        "of my knowledge and reflects the organization's compliance posture as of the "
        "report date."
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Approver: ").bold = True
    p.add_run("{{ approverName }}, {{ approverTitle }}")

    doc.add_paragraph("Signature: ________________________________  Date: __________")

    doc.save(path)


def render_via_api(template_path: Path, payload: dict, output_pdf: Path) -> dict:
    with template_path.open("rb") as fh:
        template_b64 = base64.b64encode(fh.read()).decode("utf-8")

    response = requests.post(
        f"{HOST}/document-generation/api/GenerateDocumentBase64",
        headers={
            "client_id": CLIENT_ID,
            "client_secret": CLIENT_SECRET,
            "Content-Type": "application/json",
        },
        json={
            "base64FileString": template_b64,
            "documentValues": payload,
            "outputFormat": "pdf",
        },
        timeout=60,
    )
    print(f"  HTTP {response.status_code} for {template_path.name}")
    response.raise_for_status()
    body = response.json()

    if "base64FileString" not in body:
        raise RuntimeError(f"Unexpected API response: {body}")

    pdf_bytes = base64.b64decode(body["base64FileString"])
    output_pdf.write_bytes(pdf_bytes)
    print(f"  wrote {output_pdf.name} ({len(pdf_bytes):,} bytes)")
    if not pdf_bytes.startswith(b"%PDF-"):
        raise RuntimeError(f"{output_pdf.name} is not a valid PDF")
    return body


SIMPLE_PAYLOAD = {
    "companyName": "Meridian Financial Group",
    "invoiceNumber": "INV-00471",
    "invoiceDate": "2024-01-15",
    "totalDue": 2500.00,
}

TABLE_PAYLOAD = {
    "companyName": "Meridian Financial Group",
    "invoiceDate": "2024-01-15",
    "invoiceNumber": "INV-00471",
    "lineItems": [
        {
            "description": "API Integration Consulting",
            "qty": 10,
            "unitPrice": 150.00,
            "lineTotal": 1500.00,
        },
        {
            "description": "Compliance Review",
            "qty": 5,
            "unitPrice": 200.00,
            "lineTotal": 1000.00,
        },
    ],
    "totalDue": 2500.00,
}

# Mirrors article1-client.md document_values exactly. If you change a key here,
# you MUST also update the article's Python sample, or the rendered PDF will
# have empty fields where readers expect populated data.
FULL_PAYLOAD = {
    "customer_name": "Acme Corporation",
    "invoice_number": "INV-2025-0042",
    "invoice_date": "07/15/2025",
    "due_date": "08/14/2025",
    "line_items": [
        {
            "description": "API Integration Consulting",
            "qty": 8,
            "unit_price": 195.00,
            "total": 1560.00,
        },
        {
            "description": "Document Automation Setup",
            "qty": 1,
            "unit_price": 750.00,
            "total": 750.00,
        },
    ],
    "subtotal": 2310.00,
    "tax_rate": 0.08,
    "tax_amount": 184.80,
    "total_due": 2494.80,
}


CONTRACT_STANDARD_PAYLOAD = {
    "party_a_name": "Meridian Consulting LLC",
    "party_a_address_line1": "1200 Tech Parkway, Suite 400",
    "party_a_address_city": "Austin",
    "party_a_address_state": "TX",
    "party_a_signatory_name": "Jordan Lee",
    "party_a_signatory_title": "Chief Executive Officer",
    "party_b_name": "Acme Corporation",
    "party_b_address_line1": "400 Industrial Way",
    "party_b_address_city": "Denver",
    "party_b_address_state": "CO",
    "party_b_signatory_name": "Alex Patel",
    "party_b_signatory_title": "VP Procurement",
    "effective_date": "2024-02-01",
    "term_months": 12,
    "governing_law": "Texas",
    "contract_value": 39600.00,
    "payment_schedule": "Monthly",
}

CONTRACT_AUTO_RENEWAL_PAYLOAD = dict(CONTRACT_STANDARD_PAYLOAD)

COMPLIANCE_PAYLOAD = {
    "organizationName": "Acme Corporation",
    "attestationPeriod": "Q1 2024",
    "reportDate": "2024-04-01",
    "preparedBy": "Compliance Team",
    "preparedByTitle": "Information Security Manager",
    "policyVersion": "v3.2",
    "controls": [
        {
            "control_id": "AC-01",
            "control_name": "Access Control Policy",
            "status": "Compliant",
            "evidenceLink": "https://docs.acme.com/evidence/AC-01",
        },
        {
            "control_id": "SC-28",
            "control_name": "Protection of Information at Rest",
            "status": "Compliant",
            "evidenceLink": "https://docs.acme.com/evidence/SC-28",
        },
        {
            "control_id": "IR-04",
            "control_name": "Incident Handling",
            "status": "In Remediation",
            "evidenceLink": "https://docs.acme.com/evidence/IR-04",
        },
    ],
    "totalControls": 3,
    "compliantCount": 2,
    "remediationCount": 1,
    "approverName": "Dana Okonkwo",
    "approverTitle": "Chief Information Security Officer",
}


def main() -> int:
    simple_docx = HERE / "invoice_simple.docx"
    table_docx = HERE / "invoice_table.docx"
    full_docx = HERE / "invoice_full.docx"
    contract_std_docx = HERE / "contract_standard.docx"
    contract_auto_docx = HERE / "contract_auto_renewal.docx"
    compliance_docx = HERE / "compliance_attestation.docx"

    simple_pdf = HERE / "invoice_simple_test.pdf"
    table_pdf = HERE / "invoice_table_test.pdf"
    full_pdf = HERE / "invoice_full_test.pdf"
    contract_std_pdf = HERE / "contract_standard_test.pdf"
    contract_auto_pdf = HERE / "contract_auto_renewal_test.pdf"
    compliance_pdf = HERE / "compliance_attestation_test.pdf"

    print("Building templates...")
    build_invoice_simple(simple_docx)
    build_invoice_table(table_docx)
    build_invoice_full(full_docx)
    build_contract_standard(contract_std_docx)
    build_contract_auto_renewal(contract_auto_docx)
    build_compliance_attestation(compliance_docx)
    for f in (simple_docx, table_docx, full_docx,
              contract_std_docx, contract_auto_docx, compliance_docx):
        print(f"  {f.name} ({f.stat().st_size:,} bytes)")

    print("\nRendering invoice_simple.docx via Foxit DocGen API...")
    render_via_api(simple_docx, SIMPLE_PAYLOAD, simple_pdf)

    print("\nRendering invoice_table.docx via Foxit DocGen API...")
    render_via_api(table_docx, TABLE_PAYLOAD, table_pdf)

    print("\nRendering invoice_full.docx via Foxit DocGen API...")
    render_via_api(full_docx, FULL_PAYLOAD, full_pdf)

    print("\nRendering contract_standard.docx via Foxit DocGen API...")
    render_via_api(contract_std_docx, CONTRACT_STANDARD_PAYLOAD, contract_std_pdf)

    print("\nRendering contract_auto_renewal.docx via Foxit DocGen API...")
    render_via_api(contract_auto_docx, CONTRACT_AUTO_RENEWAL_PAYLOAD, contract_auto_pdf)

    print("\nRendering compliance_attestation.docx via Foxit DocGen API...")
    render_via_api(compliance_docx, COMPLIANCE_PAYLOAD, compliance_pdf)

    print("\nAll templates rendered successfully.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
